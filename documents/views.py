from django.shortcuts import render, redirect, get_object_or_404
from django.views import View
from django.utils import timezone
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.http import HttpResponse, Http404, FileResponse
from django.conf import settings
from django.contrib import messages
from django.core.files.base import ContentFile
import os
import hashlib
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.lib import colors
from bs4 import BeautifulSoup
from .models import Equipamento, DocumentoModelo, TermoResponsabilidade, ItemTermo
from users.views import StaffRequiredMixin
from .forms import ModeloDocumentoForm
from django.urls import reverse_lazy
from reportlab.pdfgen import canvas
import io
from tinymce.widgets import TinyMCE
from django.core.exceptions import PermissionDenied
from xhtml2pdf import pisa
from django.template.loader import get_template
from django.template import Context
import tempfile
import requests
import logging

# Configurar logger
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx2pdf import convert as docx2pdf_convert
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("Bibliotecas docx e docx2pdf não disponíveis. Conversão de Word para PDF não será suportada.")

# Create your views here.

# Views de Equipamento
class EquipamentoListView(LoginRequiredMixin, ListView):
    model = Equipamento
    template_name = 'documents/equipamento_list.html'
    context_object_name = 'equipamentos'

class EquipamentoDetailView(LoginRequiredMixin, DetailView):
    model = Equipamento
    template_name = 'documents/equipamento_detail.html'
    context_object_name = 'equipamento'

class EquipamentoCreateView(StaffRequiredMixin, CreateView):
    model = Equipamento
    template_name = 'documents/equipamento_form.html'
    fields = ['tipo', 'marca', 'modelo', 'numero_serie', 'descricao',
              'valor', 'data_aquisicao', 'status', 'observacoes']

    def get_success_url(self):
        return reverse_lazy('documents:equipamento_list')

class EquipamentoUpdateView(StaffRequiredMixin, UpdateView):
    model = Equipamento
    template_name = 'documents/equipamento_form.html'
    fields = ['tipo', 'marca', 'modelo', 'numero_serie', 'descricao',
              'valor', 'data_aquisicao', 'status', 'observacoes']

    def get_success_url(self):
        return reverse_lazy('documents:equipamento_list')

# Views de Documento Modelo
class DocumentoModeloListView(StaffRequiredMixin, ListView):
    model = DocumentoModelo
    template_name = 'documents/modelo_list.html'
    context_object_name = 'modelos'

class DocumentoModeloDetailView(StaffRequiredMixin, DetailView):
    model = DocumentoModelo
    template_name = 'documents/modelo_detail.html'
    context_object_name = 'modelo'

class DocumentoModeloCreateView(StaffRequiredMixin, CreateView):
    model = DocumentoModelo
    form_class = ModeloDocumentoForm  # Use o form customizado
    template_name = 'documents/modelo_form.html'
    
    def get_success_url(self):
        return reverse_lazy('documents:modelo_list')

class DocumentoModeloUpdateView(StaffRequiredMixin, UpdateView):
    model = DocumentoModelo
    form_class = ModeloDocumentoForm  # Use o form customizado
    template_name = 'documents/modelo_form.html'
    
    def get_success_url(self):
        return reverse_lazy('documents:modelo_detail', kwargs={'pk': self.object.pk})

# Views de Termo de Responsabilidade
class TermoListView(LoginRequiredMixin, ListView):
    model = TermoResponsabilidade
    template_name = 'documents/termo_list.html'
    context_object_name = 'termos'
    
    def get_queryset(self):
        if self.request.user.is_staff:
            return TermoResponsabilidade.objects.all()
        return TermoResponsabilidade.objects.filter(colaborador=self.request.user)

class TermoDetailView(LoginRequiredMixin, DetailView):
    model = TermoResponsabilidade
    template_name = 'documents/termo_detail.html'
    context_object_name = 'termo'
    slug_field = 'uuid'
    slug_url_kwarg = 'uuid'
    
    def get_queryset(self):
        if self.request.user.is_staff:
            return TermoResponsabilidade.objects.all()
        return TermoResponsabilidade.objects.filter(colaborador=self.request.user)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        termo = self.get_object()
        
        # Adicionar o conteúdo do modelo ao contexto
        if termo.modelo:
            context['conteudo_documento'] = termo.modelo.conteudo
            
        # Verificar se o usuário pode assinar
        context['pode_assinar'] = (
            termo.status == 'PENDENTE' and 
            termo.colaborador == self.request.user
        )
        
        return context

class TermoCreateView(StaffRequiredMixin, CreateView):
    model = TermoResponsabilidade
    template_name = 'documents/termo_form.html'
    fields = ['colaborador', 'modelo', 'equipamentos', 'observacoes']
    success_url = reverse_lazy('documents:termo_list')  # Redireciona para a lista de termos
    
    def form_valid(self, form):
        form.instance.status = 'PENDENTE'
        self.object = form.save(commit=False)
        self.object.save()
        form.instance.equipamentos.clear()
        
        for equipamento in form.cleaned_data['equipamentos']:
            ItemTermo.objects.create(
                termo=self.object,
                equipamento=equipamento,
                data_entrega=timezone.now(),
                estado_entrega='Novo'
            )
        
        return super().form_valid(form)

class TermoUpdateView(UserPassesTestMixin, UpdateView):
    model = TermoResponsabilidade
    fields = ['colaborador', 'modelo', 'equipamentos', 'status', 'observacoes']
    template_name = 'documents/termo_form.html'
    context_object_name = 'termo'

    def test_func(self):
        return self.request.user.is_staff  # Apenas usuários com permissão de staff podem acessar

    def get_object(self, queryset=None):
        uuid = self.kwargs.get('uuid')  
        logger.info(f'Buscando TermoResponsabilidade com uuid={uuid}')
        obj = get_object_or_404(TermoResponsabilidade, uuid=uuid)
        logger.info(f'Objeto encontrado: {obj}')
        return obj

    def get_success_url(self):
        return reverse_lazy('documents:termo_detail', kwargs={'uuid': self.object.uuid})

class TermoSignView(LoginRequiredMixin, View):
    def get(self, request, uuid):
        try:
            termo = TermoResponsabilidade.objects.get(uuid=uuid)
        except TermoResponsabilidade.DoesNotExist:
            messages.error(request, "Termo não encontrado.")
            return redirect('documents:termo_list')
        
        # Verifica se o termo já foi assinado
        if termo.status == TermoResponsabilidade.Status.ASSINADO:
            messages.warning(request, "Este termo já foi assinado.")
            return redirect('documents:termo_detail', uuid=termo.uuid)
        
        # Verifica se o usuário tem permissão para assinar
        if not termo.pode_assinar(request.user):
            messages.error(request, "Você não tem permissão para assinar este termo.")
            return redirect('documents:termo_detail', uuid=termo.uuid)
        
        # Obtém dados do colaborador para pré-preencher o formulário
        user = termo.colaborador
        dados_colaborador = {
            'nome_completo': user.get_full_name(),
            'cpf': user.cpf if hasattr(user, 'cpf') and user.cpf else '',
            'rg': user.rg if hasattr(user, 'rg') and user.rg else '',
            'endereco': user.endereco if hasattr(user, 'endereco') and user.endereco else '',
            'numero': user.numero if hasattr(user, 'numero') and user.numero else '',
            'complemento': user.complemento if hasattr(user, 'complemento') and user.complemento else '',
            'bairro': user.bairro if hasattr(user, 'bairro') and user.bairro else '',
            'cidade': user.cidade if hasattr(user, 'cidade') and user.cidade else '',
            'estado': user.estado if hasattr(user, 'estado') and user.estado else '',
            'cep': user.cep if hasattr(user, 'cep') and user.cep else '',
        }
        
        # Obtém os equipamentos associados ao termo através da relação ItemTermo
        itens_termo = ItemTermo.objects.filter(termo=termo)
        equipamentos = [item.equipamento for item in itens_termo]
        
        context = {
            'termo': termo,
            'dados_colaborador': dados_colaborador,
            'equipamentos': equipamentos,
            'itens_termo': itens_termo,
        }
        
        return render(request, 'documents/termo_sign.html', context)
    
    def post(self, request, uuid):
        try:
            termo = TermoResponsabilidade.objects.get(uuid=uuid)
        except TermoResponsabilidade.DoesNotExist:
            messages.error(request, "Termo não encontrado.")
            return redirect('documents:termo_list')
        
        # Verifica se o termo já foi assinado
        if termo.status == TermoResponsabilidade.Status.ASSINADO:
            messages.warning(request, "Este termo já foi assinado.")
            return redirect('documents:termo_detail', uuid=termo.uuid)
        
        # Verifica se o usuário tem permissão para assinar
        if not termo.pode_assinar(request.user):
            messages.error(request, "Você não tem permissão para assinar este termo.")
            return redirect('documents:termo_detail', uuid=termo.uuid)
        
        # Obtém todos os dados do formulário
        form_data = request.POST
        
        # Verifica se os campos obrigatórios foram preenchidos
        required_fields = ['cpf', 'rg', 'endereco', 'bairro', 'cidade', 'estado', 'cep']
        for field in required_fields:
            if not form_data.get(field):
                messages.error(request, f"O campo {field.title()} é obrigatório.")
                return redirect('documents:termo_sign', uuid=termo.uuid)
        
        # Atualiza os dados do colaborador
        user = termo.colaborador
        user.cpf = form_data.get('cpf')
        user.rg = form_data.get('rg')
        user.endereco = form_data.get('endereco')
        user.numero = form_data.get('numero', '')
        user.complemento = form_data.get('complemento', '')
        user.bairro = form_data.get('bairro')
        user.cidade = form_data.get('cidade')
        user.estado = form_data.get('estado')
        user.cep = form_data.get('cep')
        user.save()
        
        # Atualiza o estado dos equipamentos e vincula o usuário aos equipamentos
        # Obter equipamentos associados ao termo através da relação ItemTermo
        itens_termo = ItemTermo.objects.filter(termo=termo)
        for item in itens_termo:
            equip = item.equipamento
            # Atribuir o usuário ao equipamento
            equip.usuario = user
            equip.status = 'EM_USO'  # Atualiza o status do equipamento para Em Uso
            equip.save()
            
            # Atualiza o estado de entrega se tiver sido informado
            estado_key = f"estado_{equip.id}"
            if estado_key in form_data:
                item.estado_entrega = form_data.get(estado_key)
                item.save()
        
        # Processa a assinatura
        termo.status = TermoResponsabilidade.Status.ASSINADO
        termo.data_assinatura = timezone.now()
        
        # Capturar o IP do cliente de forma mais completa
        # Tenta vários cabeçalhos HTTP para identificar o IP real do cliente
        cliente_ip = None
        # Verifica se está atrás de um proxy
        if 'HTTP_X_FORWARDED_FOR' in request.META:
            cliente_ip = request.META['HTTP_X_FORWARDED_FOR'].split(',')[0].strip()
        elif 'HTTP_X_REAL_IP' in request.META:
            cliente_ip = request.META['HTTP_X_REAL_IP']
        elif 'HTTP_CLIENT_IP' in request.META:
            cliente_ip = request.META['HTTP_CLIENT_IP']
        elif 'REMOTE_ADDR' in request.META:
            cliente_ip = request.META['REMOTE_ADDR']
        
        # Captura informações do navegador/dispositivo
        user_agent = request.META.get('HTTP_USER_AGENT', 'Navegador não identificado')
        
        # Usa o IP capturado ou "não identificado" se não encontrado
        termo.ip_assinatura = cliente_ip or "IP não identificado"
        
        # Adiciona informações do dispositivo usado para assinatura
        termo.observacoes += f"\n\nInformações da assinatura:\nIP: {termo.ip_assinatura}\nDispositivo: {user_agent}"
        
        # Gera hash para a assinatura
        signature_data = f"{termo.uuid}_{termo.colaborador.username}_{termo.data_assinatura.isoformat()}_{termo.ip_assinatura}"
        termo.hash_assinatura = hashlib.sha256(signature_data.encode()).hexdigest()
        
        termo.save()
        
        # Gerar o PDF do termo
        self.gerar_pdf_termo(termo)
        
        messages.success(request, "Termo assinado com sucesso!")
        return redirect('documents:termo_detail', uuid=termo.uuid)
    
    def gerar_pdf_termo(self, termo, custom_path=None):
        """
        Gera o PDF do termo de responsabilidade assinado usando uma abordagem mais simples
        
        Args:
            termo: Objeto TermoResponsabilidade
            custom_path: Caminho personalizado para o arquivo PDF (opcional)
        """
        logger.info(f"Iniciando geração de PDF para termo {termo.uuid}")
        
        # Verificar o modelo
        if not termo.modelo:
            logger.error(f"Termo {termo.uuid} não possui modelo associado")
            raise Exception("O termo não possui um modelo de documento associado.")

        # Extrair o user_agent das observações, se existir
        user_agent = None
        if termo.observacoes and 'Dispositivo:' in termo.observacoes:
            try:
                for linha in termo.observacoes.split('\n'):
                    if linha.strip().startswith('Dispositivo:'):
                        user_agent = linha.replace('Dispositivo:', '').strip()
                        break
            except Exception as e:
                logger.error(f"Erro ao extrair user_agent: {e}")
        
        # Configuração do caminho do PDF
        if custom_path:
            pdf_path = custom_path
        else:
            media_root = settings.MEDIA_ROOT
            pdf_dir = os.path.join(media_root, 'documentos', 'termos')
            os.makedirs(pdf_dir, exist_ok=True)
            file_name = f"termo_{termo.uuid}.pdf"
            pdf_path = os.path.join(pdf_dir, file_name)
        
        # Verificar se existe um arquivo Word associado ao modelo e se as bibliotecas necessárias estão disponíveis
        if DOCX_AVAILABLE and termo.modelo and termo.modelo.arquivo_word:
            logger.info(f"Gerando PDF a partir do arquivo Word: {termo.modelo.arquivo_word.path}")
            
            # Verificar se o arquivo Word existe
            if not os.path.exists(termo.modelo.arquivo_word.path):
                logger.error(f"Arquivo Word não encontrado: {termo.modelo.arquivo_word.path}")
                raise Exception(f"Arquivo do modelo não encontrado: {termo.modelo.arquivo_word.path}")
                
            try:
                # Criar um arquivo temporário para trabalhar com o Word
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                    temp_docx_path = temp_docx.name
                    with open(termo.modelo.arquivo_word.path, 'rb') as original_file:
                        temp_docx.write(original_file.read())
                
                # Substituir placeholders no documento Word
                doc = Document(temp_docx_path)
                user = termo.colaborador
                nome_completo = f"{user.first_name} {user.last_name}"
                
                # Obter o endereço completo
                endereco_completo = f"{user.endereco}, {user.numero}"
                if user.complemento:
                    endereco_completo += f", {user.complemento}"
                endereco_completo += f", {user.bairro}, {user.cidade}/{user.estado}, CEP: {user.cep}"
                
                # Placeholders para substituir
                placeholders = {
                    "${NOME}": nome_completo,
                    "${CPF}": user.cpf or "",
                    "${RG}": user.rg or "",
                    "${ENDERECO}": user.endereco or "",
                    "${NUMERO}": user.numero or "",
                    "${COMPLEMENTO}": user.complemento or "",
                    "${BAIRRO}": user.bairro or "",
                    "${CIDADE}": user.cidade or "",
                    "${ESTADO}": user.estado or "",
                    "${CEP}": user.cep or "",
                    "${ENDERECO_COMPLETO}": endereco_completo,
                    "${DATA_ASSINATURA}": termo.data_assinatura.strftime("%d/%m/%Y") if termo.data_assinatura else "",
                    "${IP_ASSINATURA}": termo.ip_assinatura or "",
                    "${DISPOSITIVO_ASSINATURA}": user_agent if 'user_agent' in locals() else "",
                    "${HASH_ASSINATURA}": termo.hash_assinatura or "",
                }
                
                # Adicionar informações dos equipamentos para a tabela
                itens_termo = ItemTermo.objects.filter(termo=termo)
                equipamentos = [item.equipamento for item in itens_termo]
                
                # Para cada equipamento, adicione placeholders específicos
                for i, equip in enumerate(equipamentos):
                    placeholders[f"${{EQUIP_{i+1}_DESCRICAO}}"] = f"{equip.tipo} {equip.marca} {equip.modelo} - {equip.numero_serie}"
                    placeholders[f"${{EQUIP_{i+1}_VALOR}}"] = f"R$ {equip.valor:.2f}".replace('.', ',')
                
                # Placeholder para valor total
                total_valor = sum(equip.valor for equip in equipamentos)
                placeholders["${VALOR_TOTAL}"] = f"R$ {total_valor:.2f}".replace('.', ',')
                
                # Procurar e substituir texto em todos os parágrafos
                for paragraph in doc.paragraphs:
                    for key, value in placeholders.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value or ""))
                
                # Procurar e substituir texto em todas as tabelas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for key, value in placeholders.items():
                                    if key in paragraph.text:
                                        paragraph.text = paragraph.text.replace(key, str(value or ""))
                
                # Salvar o documento com as substituições
                doc.save(temp_docx_path)
                
                # Converter o arquivo docx para PDF
                docx2pdf_convert(temp_docx_path, pdf_path)
                
                # Limpar o arquivo temporário
                os.unlink(temp_docx_path)
                
                # Atualizar o termo com o caminho do PDF apenas se não for uma prévia
                if not custom_path:
                    termo.arquivo_pdf.save(file_name, ContentFile(open(pdf_path, 'rb').read()), save=True)
                
                return pdf_path
                
            except Exception as e:
                # Se ocorrer um erro, voltar para o método padrão de geração de PDF
                logger.error(f"Erro ao converter Word para PDF: {str(e)}")
        
        # Se não tiver arquivo Word ou ocorrer erro, usar a abordagem HTML
        # Busca dados do termo
        user = termo.colaborador
        itens_termo = ItemTermo.objects.filter(termo=termo)
        equipamentos = [item.equipamento for item in itens_termo]
        
        # Informações do colaborador
        nome_completo = f"{user.first_name} {user.last_name}"
        
        # Obter o endereço completo
        endereco_completo = f"{user.endereco}, {user.numero}"
        if user.complemento:
            endereco_completo += f", {user.complemento}"
        endereco_completo += f", {user.bairro}, {user.cidade}/{user.estado}, CEP: {user.cep}"
        
        # Detalhes dos equipamentos
        lista_equipamentos = []
        total_valor = 0
        for i, equip in enumerate(equipamentos):
            item = itens_termo[i]
            lista_equipamentos.append({
                'numero': i + 1,
                'tipo': equip.tipo,
                'marca': equip.marca,
                'modelo': equip.modelo,
                'numero_serie': equip.numero_serie,
                'descricao': f"{equip.tipo} {equip.marca} {equip.modelo} - {equip.numero_serie}",
                'estado': item.estado_entrega,
                'valor': equip.valor,
                'valor_formatado': f"R$ {equip.valor:.2f}".replace('.', ',')
            })
            total_valor += equip.valor
        
        # Data de assinatura formatada
        data_assinatura = termo.data_assinatura.strftime("%d/%m/%Y") if termo.data_assinatura else ""
        
        # Conteúdo do termo (texto limpo)
        conteudo_termo = ""
        if termo.modelo and termo.modelo.conteudo:
            soup = BeautifulSoup(termo.modelo.conteudo, 'html.parser')
            conteudo_termo = soup.get_text('\n\n')
        
        # Usar template HTML para gerar o PDF
        context = {
            'titulo': "TERMO DE RESPONSABILIDADE",
            'nome_colaborador': nome_completo,
            'cpf': user.cpf,
            'rg': user.rg,
            'endereco': user.endereco,
            'numero': user.numero,
            'complemento': user.complemento,
            'bairro': user.bairro,
            'cidade': user.cidade,
            'estado': user.estado,
            'cep': user.cep,
            'endereco_completo': endereco_completo,
            'equipamentos': lista_equipamentos,
            'total_valor': total_valor,
            'total_valor_formatado': f"R$ {total_valor:.2f}".replace('.', ','),
            'conteudo_termo': conteudo_termo,
            'data_assinatura': data_assinatura,
            'ip_assinatura': termo.ip_assinatura,
            'dispositivo_assinatura': user_agent,
            'hash_assinatura': termo.hash_assinatura
        }
        
        # Renderizar o template HTML
        template = get_template('documents/termo_pdf_template.html')
        html = template.render(context)
        
        # Converter HTML para PDF
        with open(pdf_path, 'wb') as pdf_file:
            pisa_status = pisa.CreatePDF(html, dest=pdf_file)
            
        # Verificar se a conversão foi bem-sucedida
        if pisa_status.err:
            raise Exception("Erro ao gerar o PDF")
            
        # Atualizar o termo com o caminho do PDF apenas se não for uma prévia
        if not custom_path:
            termo.arquivo_pdf.save(file_name, ContentFile(open(pdf_path, 'rb').read()), save=True)
        
        return pdf_path

class TermoDownloadView(LoginRequiredMixin, View):
    def get(self, request, uuid):
        try:
            termo = TermoResponsabilidade.objects.get(uuid=uuid)
        except TermoResponsabilidade.DoesNotExist:
            messages.error(request, "Termo não encontrado.")
            return redirect('documents:termo_list')
        
        # Verifica se o termo já foi assinado
        if termo.status != TermoResponsabilidade.Status.ASSINADO:
            messages.warning(request, "Este termo ainda não foi assinado e não possui PDF para download.")
            return redirect('documents:termo_detail', uuid=termo.uuid)
        
        # Verifica se o usuário tem permissão para baixar
        # O usuário deve ser o colaborador associado ao termo ou um administrador
        if request.user != termo.colaborador and not request.user.is_staff:
            messages.error(request, "Você não tem permissão para baixar este documento.")
            return redirect('documents:termo_detail', uuid=termo.uuid)
        
        # Verifica se o arquivo PDF existe
        if not termo.arquivo_pdf:
            # Tenta gerar o PDF caso ele não exista
            try:
                # Instanciar TermoSignView para gerar o PDF
                sign_view = TermoSignView()
                sign_view.gerar_pdf_termo(termo)
                # Recarrega o termo para obter o arquivo atualizado
                termo.refresh_from_db()
                
                if not termo.arquivo_pdf:
                    messages.error(request, "Não foi possível gerar o PDF do termo.")
                    return redirect('documents:termo_detail', uuid=termo.uuid)
            except Exception as e:
                messages.error(request, f"Erro ao gerar o PDF: {str(e)}")
                return redirect('documents:termo_detail', uuid=termo.uuid)
        
        # Verifica se o arquivo físico existe
        try:
            arquivo_path = termo.arquivo_pdf.path
            
            if not os.path.exists(arquivo_path):
                messages.error(request, "O arquivo PDF não foi encontrado no servidor.")
                return redirect('documents:termo_detail', uuid=termo.uuid)
            
            # Define o nome do arquivo para download
            nome_colaborador = termo.colaborador.get_full_name().replace(' ', '_')
            nome_arquivo = f"termo_{nome_colaborador}_{termo.uuid}.pdf"
            
            # Retorna o arquivo como resposta
            with open(arquivo_path, 'rb') as pdf:
                if 'download' in request.GET:
                    # Se o parâmetro download estiver presente, força o download
                    response = HttpResponse(pdf.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
                else:
                    # Caso contrário, exibe no navegador
                    response = HttpResponse(pdf.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'inline; filename="{nome_arquivo}"'
                    response['X-Frame-Options'] = 'SAMEORIGIN'
                    response['Content-Security-Policy'] = "frame-ancestors 'self'"
                
                response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
                response['Pragma'] = 'no-cache'
                response['Expires'] = '0'
                return response
            
        except Exception as e:
            messages.error(request, f"Erro ao acessar o arquivo: {str(e)}")
            return redirect('documents:termo_detail', uuid=termo.uuid)

class TermoPreviewView(LoginRequiredMixin, View):
    def get(self, request, uuid):
        logger.info(f"Iniciando preview do termo {uuid}")
        
        try:
            termo = TermoResponsabilidade.objects.get(uuid=uuid)
            logger.info(f"Termo encontrado: {termo.id}, modelo: {termo.modelo_id if termo.modelo else 'Nenhum'}")
        except TermoResponsabilidade.DoesNotExist:
            logger.error(f"Termo não encontrado: {uuid}")
            messages.error(request, "Termo não encontrado.")
            return redirect('documents:termo_list')
        
        # Verifica se o usuário tem permissão para visualizar
        if not termo.pode_assinar(request.user) and not request.user.is_staff:
            logger.warning(f"Usuário {request.user.username} sem permissão para visualizar termo {termo.uuid}")
            messages.error(request, "Você não tem permissão para visualizar este termo.")
            return redirect('documents:termo_list')
        
        # Preparar dados do colaborador para prévia
        user = termo.colaborador
        
        # Se os dados existem, use-os; caso contrário, use dados do formulário se disponíveis
        form_data = request.session.get(f'termo_preview_data_{uuid}', {})
        
        # Verificar se o modelo existe
        if not termo.modelo:
            logger.error(f"Termo {termo.uuid} não possui modelo de documento associado")
            messages.error(request, "Este termo não possui um modelo de documento associado.")
            return redirect('documents:termo_detail', uuid=termo.uuid)
            
        # Criar uma versão temporária do PDF para prévia
        try:
            logger.info(f"Gerando prévia do PDF para o termo {termo.uuid}")
            
            # Criar uma instância da TermoSignView para usar o método gerar_pdf_termo
            sign_view = TermoSignView()
            
            # Simular temporariamente alguns dados para a prévia
            # Salvar valores originais
            original_ip = termo.ip_assinatura
            original_hash = termo.hash_assinatura
            original_status = termo.status
            original_data = termo.data_assinatura
            
            # Temporariamente definir valores para a prévia
            termo.ip_assinatura = "Prévia - Não assinado"
            termo.hash_assinatura = "Prévia - Este documento ainda não foi assinado"
            termo.status = TermoResponsabilidade.Status.PENDENTE
            termo.data_assinatura = timezone.now()  # Data atual apenas para a prévia
            
            # Gerar o PDF temporário
            preview_dir = os.path.join(settings.MEDIA_ROOT, 'documentos', 'previews')
            os.makedirs(preview_dir, exist_ok=True)
            
            preview_filename = f"preview_{termo.uuid}_{int(timezone.now().timestamp())}.pdf"
            preview_path = os.path.join(preview_dir, preview_filename)
            
            # Usa a função existente para gerar o PDF, mas com um caminho diferente
            try:
                pdf_path = sign_view.gerar_pdf_termo(termo, custom_path=preview_path)
                logger.info(f"Prévia do PDF gerada com sucesso: {pdf_path}")
            except Exception as e:
                logger.error(f"Erro ao gerar PDF na prévia: {str(e)}")
                raise
            
            # Restaurar valores originais
            termo.ip_assinatura = original_ip
            termo.hash_assinatura = original_hash
            termo.status = original_status
            termo.data_assinatura = original_data
            termo.save()  # Salvar para garantir que não foram feitas alterações permanentes
            
            # Retornar o PDF como resposta
            if os.path.exists(preview_path):
                try:
                    with open(preview_path, 'rb') as pdf:
                        file_content = pdf.read()
                        
                    if not file_content:
                        logger.error(f"Arquivo de prévia vazio: {preview_path}")
                        messages.error(request, "Arquivo de prévia vazio. Tente novamente.")
                        return redirect('documents:termo_detail', uuid=termo.uuid)
                        
                    response = HttpResponse(file_content, content_type='application/pdf')
                    response['Content-Disposition'] = f'inline; filename="preview_{termo.uuid}.pdf"'
                    response['X-Frame-Options'] = 'SAMEORIGIN'
                    response['Content-Security-Policy'] = "frame-ancestors 'self'"
                    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
                    response['Pragma'] = 'no-cache'
                    response['Expires'] = '0'
                    logger.info(f"Arquivo PDF sendo servido: {preview_path}")
                    return response
                except Exception as e:
                    logger.error(f"Erro ao ler o arquivo de prévia: {str(e)}")
                    messages.error(request, f"Erro ao ler o arquivo de prévia: {str(e)}")
                    return redirect('documents:termo_detail', uuid=termo.uuid)
            else:
                logger.error(f"Arquivo de prévia não foi encontrado: {preview_path}")
                messages.error(request, "Não foi possível gerar a prévia do documento.")
                return redirect('documents:termo_detail', uuid=termo.uuid)
            
        except Exception as e:
            logger.error(f"Erro ao gerar prévia do documento: {e}")
            messages.error(request, f"Erro ao gerar prévia: {str(e)}")
            return redirect('documents:termo_sign', uuid=termo.uuid)